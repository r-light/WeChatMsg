# analysis 群聊聊天记录分析

import re
import os
import sys
import time
import codecs
import json
import jieba
import random
import datetime
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from wordcloud import WordCloud
from collections import Counter
from matplotlib.font_manager import FontProperties
from docx import Document

# 设置全局字体
plt.rcParams['font.sans-serif'] = ['SimHei']  # 替换为你系统中支持的中文字体，这里使用的是宋体
plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号+

# 设置为使用非交互式后端（Agg）
plt.switch_backend('agg')

    # 例如，获取当前目录下的stopwords文件夹中的所有txt文件的内容
stopwords_directory = './stopwords'
stopwords=[]


def load_stopwords(directory_path):
    stopwords = []

    try:
        # 获取目录下所有txt文件的文件名
        txt_files = [f for f in os.listdir(directory_path) if f.endswith('.txt')]

        # 读取每个txt文件的内容并合并
        for txt_file in txt_files:
            file_path = os.path.join(directory_path, txt_file)
            with codecs.open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                stopwords.extend(file.read().splitlines())

    except OSError as e:
        print(f"Error reading directory '{directory_path}': {e}")

    return stopwords

# 使用别人的stopwords
# stopwords = load_stopwords(stopwords_directory)
# 使用自己的stopwords
stopwords=[
    '的', '是', '在', '了', '和', '与',
    '而', '但', '因此', '然而', '虽然', '所以',
    '我', '你', '他', '它', '这', '那','们','就','个','吗','不','都','还','也'
    '现在', '时候', '今天', '明天', '之前', '之后',
    '这里', '那里', '这儿', '那儿',
    '一', '二', '三', '十', '百', '千'
]

# 读取数据
def read_data(file_name):
    with open(file_name, 'r', encoding='utf-8') as f:
        html_content = f.read()


    pattern = r'([a-zA-Z_]\w*):'
    # 找到所有的匹配模式的单词，并去重
    matches = list(set(re.findall(pattern, html_content)))
    # 在pattern两边加上双引号
    result = re.sub(pattern, r'"\1":', html_content)
    # 使用正则表达式将单引号替换为双引号
    result = re.sub(r"'", '"', result)

    # 定义你想要添加单引号的关键词列表
    keywords = ['type', 'text', 'is_send', 'avatar_path', 'timestamp', 'is_chatroom', 'displayname','refer_text','sub_type']

    # 对于在matches但是不在keywords列表中的单词，进行替换
    for match in matches:
        if match not in keywords:
            result = result.replace(f'"{match}"', match)

    data_json = json.loads(result)
    # 转换为Pandas DataFrame
    df = pd.DataFrame(data_json)

    return df


# 数据清洗
def data_clean(data):
    # 清理数据：填充不存在的字段
    data = data.fillna('')
    # 去除重复数据
    data.drop_duplicates(inplace=True)
    # 去除缺失值
    data.dropna(inplace=True)
    # 去除无用列
    data.drop([ 'is_send', 'avatar_path', 'is_chatroom', 'refer_text', 'sub_type'], axis=1, inplace=True)
    # 重置索引
    data.reset_index(drop=True, inplace=True)
    return data


# 数据预处理
def data_preprocessing(data):
    # 只保留 type 为 1 和 49 的记录
    data = data[(data['type'] == 1) | (data['type'] == 49)]
    
     # Convert 'timestamp' to datetime with UTC
    data['timestamp'] = pd.to_datetime(data['timestamp'], unit='s', utc=True)
    # 将UTC时间转换为北京时间 (UTC+8)
    data['timestamp'] = data['timestamp'].dt.tz_convert('Asia/Shanghai')
    
    # 添加时间特征列
    data['year'] = data['timestamp'].dt.year
    data['month'] = data['timestamp'].dt.month
    data['date'] = data['timestamp'].dt.date
    data['time'] = data['timestamp'].dt.time
    data['weekday'] = data['timestamp'].dt.weekday + 1  # 从1开始，周一为1，周日为7
    data['hour'] = data['timestamp'].dt.hour
    # data['minute'] = data['timestamp'].dt.minute
    # data['second'] = data['timestamp'].dt.second
    
    return data



# 创建word 文件
def create_word_file():
    # 创建word文档
    document = Document()
    # 添加标题
    document.add_heading('群聊聊天记录分析', 0)
    # 添加段落
    document.add_paragraph('本文档是对群聊聊天记录的分析结果。')
    # 添加分页符
    document.add_page_break()
    # 保存文档
    document.save('群聊聊天记录分析.docx')
    print('word文件创建成功！')
    return document

# 群聊元数据
def group_metadata(data, document):
    # 统计群聊人数
    group_num = len(data['displayname'].unique())
    # 统计群聊起止日期
    group_start_date = data['date'].min()
    group_end_date = data['date'].max()
    # 统计群聊天数
    group_date_num = len(data['date'].unique())
    # 统计群聊天记录数
    group_chat_num = len(data)
    # 统计群聊天记录字数
    data['text_length'] = data['text'].apply(lambda x: len(x))
    group_chat_length = data['text_length'].sum()

   
    # 添加段落
    document.add_paragraph('群聊人数：{}'.format(group_num))
    document.add_paragraph('群聊起止日期：{} - {}'.format(group_start_date, group_end_date))
    document.add_paragraph('群聊天数：{}'.format(group_date_num))
    document.add_paragraph('群聊天记录数：{}'.format(group_chat_num))
    document.add_paragraph('群聊天记录字数：{}'.format(group_chat_length))
    # 添加分页符
    document.add_page_break()



 # 群聊发言词云分析，群聊发言词云
def word_cloud_analysis(data, document):
    # 统计群聊发言词频
    text_data = data['text'].str.cat(sep=' ')
    words = jieba.cut(text_data, cut_all=False)
    word_list = ' '.join(words)

   
    for stopword in stopwords:
        word_list = word_list.replace(stopword, '')

    # 绘制词云
    wordcloud = WordCloud(
        font_path='C:/Windows/Fonts/simfang.ttf',  # 使用中文字体，确保字体文件存在
        background_color='white',  # 背景颜色
        width=800, height=400      # 图像大小
    ).generate(word_list)
   
    # 保存图片
    wordcloud.to_file('群聊发言词云.png')
    # 添加图片
    document.add_picture('群聊发言词云.png')
    # 添加分页符
    document.add_page_break()

# 每个人的发言词云
def word_cloud_analysis_by_user(data, document):
    for displayname in data['displayname'].unique():
        # 添加二级标题
        document.add_heading(displayname, level=2)

        filtered_df = data[data['displayname'] == displayname]
        text_data = filtered_df['text'].str.cat(sep=' ')
        words = jieba.cut(text_data, cut_all=False)
        word_list = ' '.join(words)


        for stopword in stopwords:
            word_list = word_list.replace(stopword, '')


        wordcloud = WordCloud(
            font_path='C:/Windows/Fonts/simfang.ttf',  # 使用中文字体，确保字体文件存在
            background_color='white',  # 背景颜色
            width=800, height=400      # 图像大小
        ).generate(word_list)

         # 保存图片
        wordcloud.to_file('{}_发言词云.png'.format(displayname))
        # 添加图片
        document.add_picture('{}_发言词云.png'.format(displayname))
        # 添加分页符
        document.add_page_break()


# 统计每个人的发言次数
def count_text_times_by_user(data, document):
    # 统计每个人的发言次数
    data_count = data.groupby('displayname')['displayname'].count()
    # 统计每个人的发言次数并排序
    data_count = data_count.sort_values(ascending=False)
    # 绘制柱状图
    data_count.plot(kind='bar', figsize=(10, 5), title='发言次数统计')
    # 保存图片
    plt.savefig('发言次数统计.png')
    # 添加图片
    document.add_picture('发言次数统计.png')
    # 添加分页符
    document.add_page_break()

# 统计每个人的发言字数
def count_text_counts_by_user(data, document):
    # 统计每个人的发言字数
    data['text_length'] = data['text'].apply(lambda x: len(x))
    data_length = data.groupby('displayname')['text_length'].sum()
    # 统计每个人的发言字数并排序
    data_length = data_length.sort_values(ascending=False)
    # 绘制柱状图
    data_length.plot(kind='bar', figsize=(10, 5), title='发言字数统计')
    # 保存图片
    plt.savefig('发言字数统计.png')
    # 添加图片
    document.add_picture('发言字数统计.png')
    # 添加分页符
    document.add_page_break()


def count_text_stat_by_user(data, document, user_name, stat_type):
    # 统计 user_name 的发言时间/日期/星期/月份的分布
    filtered_df = data[data['displayname'] == user_name]

    if stat_type == 'time':
        # 统计发言时间的分布
        stat_counts = filtered_df['timestamp'].dt.hour.value_counts().sort_index()
        x_label = '小时'
        title = '发言时间统计'
    elif stat_type == 'date':
        # 统计发言日期的分布
        stat_counts = filtered_df['date'].value_counts().sort_index()
        x_label = '日期'
        title = '发言日期统计'
    elif stat_type == 'weekday':
        # 统计发言星期的分布
        stat_counts = filtered_df['weekday'].value_counts().sort_index()
        x_label = '星期'
        title = '发言星期统计'
    elif stat_type == 'month':
        # 统计发言月份的分布
        stat_counts = filtered_df['month'].value_counts().sort_index()
        x_label = '月份'
        title = '发言月份统计'
    else:
        raise ValueError("Invalid stat_type. Supported values: 'time', 'date', 'weekday', 'month'")

    # 创建新的 Figure 和 Axes 对象
    fig, ax = plt.subplots()

    # 画柱状图
    ax.bar(stat_counts.index, stat_counts.values, align='center', alpha=0.7)
    ax.set_xlabel(x_label)
    ax.set_ylabel('发言次数')
    ax.set_title(f'{user_name} {title}')

    # 保存图片
    plt.savefig(f'{user_name}_{title}.png')

    # 关闭当前图形
    plt.close(fig)

    # 添加图片到文档
    document.add_picture(f'{user_name}_{title}.png')

    # 添加分页符
    document.add_page_break()



# 数据分析并写入word文档
def data_analysis(data, document):

   
    # 群聊元数据分析，群聊人数，群聊起止日期，群聊天数，群聊天记录数，群聊天记录字数
    # 添加标题
    document.add_heading('群聊元数据分析', level=1)
    group_metadata(data, document)

    # 群聊发言次数统计，每个人的发言次数，每个人的发言字数，每个人的发言时间，每个人的发言日期，每个人的发言星期，每个人的发言月份

    # 统计每个人的发言次数
    # 添加标题
    document.add_heading('群聊发言次数统计', level=1)
    count_text_times_by_user(data, document)
    # 统计每个人的发言字数
    # 添加标题
    document.add_heading('群聊发言字数统计', level=1)
    count_text_counts_by_user(data, document)
    # 添加标题
    document.add_heading('每个人的聊天时间分布统计', level=1)
    # 获取 displayname 列表
    displayname_list = data['displayname'].unique()
    # 遍历 displayname 列表
    for displayname in displayname_list:
        # 添加二级标题
        document.add_heading('{}'.format(displayname), level=2)
        # 统计每个人的发言时间,发言日期,发言星期,发言月份
        count_text_stat_by_user(data, document, displayname, 'time')
        count_text_stat_by_user(data, document, displayname, 'date')
        count_text_stat_by_user(data, document, displayname, 'weekday')
        count_text_stat_by_user(data, document, displayname, 'month')


    # 群聊发言词云分析，群聊发言词云
    # 添加标题
    document.add_heading('群聊词云', level=1)
    word_cloud_analysis(data, document)
    # 每个人的发言词云
    # 添加标题
    document.add_heading('每个人的发言词云', level=1)
    word_cloud_analysis_by_user(data, document)

    # 保存文档
    document.save('群聊聊天记录分析.docx')


    
if __name__ == '__main__':
    # 读取数据
    file_path = 'test.txt' 

    data = read_data(file_path)
    # 数据清洗
    data = data_clean(data)
    # 数据预处理
    data = data_preprocessing(data)
    # 创建word文档
    document = create_word_file()
    # 数据分析并写入word文档
    data_analysis(data, document)
    print('数据分析完成！')